import streamlit as st
import pandas as pd
import mysql.connector
from io import BytesIO
from docx import Document
from fpdf import FPDF
import os

# ------------------- CONFIGURAÇÃO -------------------
st.set_page_config(page_title="Relatórios - Veículos Elétricos", layout="wide")
st.title("📊 Relatórios de Veículos Elétricos")
st.sidebar.header("Filtros de Pesquisa")

# ------------------- CONEXÃO -------------------
@st.cache_resource
def get_conn():
    caminho_ca = os.path.join(os.path.dirname(__file__), st.secrets["SSL_CA_PATH"])
    return mysql.connector.connect(
        host=st.secrets["MYSQL_HOST"],
        user=st.secrets["MYSQL_USER"],
        password=st.secrets["MYSQL_PASSWORD"],
        database=st.secrets["MYSQL_DATABASE"],
        port=st.secrets["MYSQL_PORT"],
        ssl_ca=caminho_ca,
        ssl_verify_cert=True
    )

conn = get_conn()

# ------------------- FUNÇÃO AUXILIAR -------------------
def adicionar_opcao_todos(lista):
    return ["Todos"] + lista

def sql_in_list(int_list):
    return "(" + ",".join(str(x) for x in int_list) + ")"

# ------------------- FILTROS -------------------
# Região
regioes_df = pd.read_sql("SELECT DISTINCT regiao, id_regiao FROM regiao ORDER BY regiao", conn)
regioes = adicionar_opcao_todos(regioes_df['regiao'].tolist())
regioes_sel = st.sidebar.multiselect("🌎 Regiões", regioes)

# Estados
if regioes_sel and "Todos" not in regioes_sel:
    ids_regioes = regioes_df[regioes_df['regiao'].isin(regioes_sel)]['id_regiao'].tolist()
    estados_df = pd.read_sql(f"SELECT uf, id_estado FROM estado WHERE id_regiao IN {sql_in_list(ids_regioes)} ORDER BY uf", conn)
else:
    estados_df = pd.read_sql("SELECT uf, id_estado FROM estado ORDER BY uf", conn)
estados = adicionar_opcao_todos(estados_df['uf'].tolist())
estados_sel = st.sidebar.multiselect("🏙️ Estados", estados)

# Cidades
if estados_sel and "Todos" not in estados_sel:
    ids_estados = estados_df[estados_df['uf'].isin(estados_sel)]['id_estado'].tolist()
    cidades_df = pd.read_sql(f"SELECT cidade, id_cidade FROM cidade WHERE id_estado IN {sql_in_list(ids_estados)} ORDER BY cidade", conn)
else:
    cidades_df = pd.read_sql("SELECT cidade, id_cidade FROM cidade ORDER BY cidade", conn)
cidades = adicionar_opcao_todos(cidades_df['cidade'].tolist())
cidades_sel = st.sidebar.multiselect("🌆 Cidades", cidades)

# Ano
anos_df = pd.read_sql("SELECT DISTINCT ano_inicial AS ano FROM cidade_tipo_modelo ORDER BY ano_inicial", conn)
anos = adicionar_opcao_todos(anos_df['ano'].astype(str).tolist())
anos_sel = st.sidebar.multiselect("📅 Ano", anos)

# Tecnologia
tecnologias_df = pd.read_sql("SELECT DISTINCT tecnologia, id_tecnologia FROM tecnologia ORDER BY tecnologia", conn)
tecnologias = adicionar_opcao_todos(tecnologias_df['tecnologia'].tolist())
tecnologias_sel = st.sidebar.multiselect("⚡ Tecnologia", tecnologias)

# Classificação
classificacao_df = pd.read_sql("SELECT DISTINCT classificacao, id_classificacao FROM classificacao_veiculo ORDER BY classificacao", conn)
classificacoes = adicionar_opcao_todos(classificacao_df['classificacao'].tolist())
classificacoes_sel = st.sidebar.multiselect("🚘 Classificação", classificacoes)

# ------------------- BOTÕES -------------------
st.sidebar.markdown("---")
col_btn1, col_btn2 = st.sidebar.columns(2)
pesquisar = col_btn1.button("Pesquisar", use_container_width=True)
limpar = col_btn2.button("Limpar Filtros", use_container_width=True)

# ------------------- LÓGICA DE LIMPEZA -------------------
if limpar:
    st.session_state.clear()
    st.experimental_rerun()

# ------------------- CONSULTA -------------------
if pesquisar:
    query = """
    SELECT 
        m.modelo,
        SUM(ctm.quantidade) AS quantidade,
        t.tecnologia,
        cv.classificacao
    FROM cidade_tipo_modelo ctm
    JOIN modelo m ON ctm.id_modelo = m.id_modelo
    JOIN tecnologia t ON ctm.id_tecnologia = t.id_tecnologia
    JOIN classificacao_veiculo cv ON ctm.id_classificacao = cv.id_classificacao
    JOIN cidade c ON ctm.id_cidade = c.id_cidade
    JOIN estado e ON c.id_estado = e.id_estado
    JOIN regiao r ON e.id_regiao = r.id_regiao
    WHERE ctm.quantidade > 0
    """

    # Aplicando filtros
    if regioes_sel and "Todos" not in regioes_sel:
        ids = regioes_df[regioes_df["regiao"].isin(regioes_sel)]["id_regiao"].tolist()
        query += f" AND r.id_regiao IN {sql_in_list(ids)}"
    if estados_sel and "Todos" not in estados_sel:
        ids = estados_df[estados_df["uf"].isin(estados_sel)]["id_estado"].tolist()
        query += f" AND e.id_estado IN {sql_in_list(ids)}"
    if cidades_sel and "Todos" not in cidades_sel:
        ids = cidades_df[cidades_df["cidade"].isin(cidades_sel)]["id_cidade"].tolist()
        query += f" AND c.id_cidade IN {sql_in_list(ids)}"
    if anos_sel and "Todos" not in anos_sel:
        anos_n = [int(a) for a in anos_sel if a != "Todos"]
        query += f" AND (ctm.ano_inicial IN {sql_in_list(anos_n)} OR ctm.ano_final IN {sql_in_list(anos_n)})"
    if tecnologias_sel and "Todos" not in tecnologias_sel:
        ids = tecnologias_df[tecnologias_df["tecnologia"].isin(tecnologias_sel)]["id_tecnologia"].tolist()
        query += f" AND t.id_tecnologia IN {sql_in_list(ids)}"
    if classificacoes_sel and "Todos" not in classificacoes_sel:
        ids = classificacao_df[classificacao_df["classificacao"].isin(classificacoes_sel)]["id_classificacao"].tolist()
        query += f" AND cv.id_classificacao IN {sql_in_list(ids)}"

    query += """
    GROUP BY m.modelo, t.tecnologia, cv.classificacao
    ORDER BY m.modelo
    """

    df = pd.read_sql(query, conn)

    if not df.empty:
        # Total de veículos
        total_veiculos = df['quantidade'].sum()

        filtros_txt = []
        for nome, valor in {
            "Região": regioes_sel, "Estado": estados_sel, "Cidade": cidades_sel,
            "Ano": anos_sel, "Tecnologia": tecnologias_sel, "Classificação": classificacoes_sel
        }.items():
            if valor:
                filtros_txt.append(f"{nome}: {', '.join(map(str, valor))}")

        st.subheader(f"🔎 Resultados para a pesquisa: {' | '.join(filtros_txt) if filtros_txt else '(sem filtros)'}")

        # KPIs
        c1, c2 = st.columns(2)
        c1.metric("Modelos únicos", df["modelo"].nunique())
        c2.metric("Total de veículos", total_veiculos)

        st.dataframe(df, use_container_width=True)

        # ------------------- BOTÕES DE DOWNLOAD -------------------
        col1, col2, col3, col4 = st.columns(4)

        # CSV
        csv = df.to_csv(index=False).encode('utf-8')
        col1.download_button("📥 CSV", data=csv, file_name='relatorio.csv', mime='text/csv')

        # Excel
        excel_buffer = BytesIO()
        df.to_excel(excel_buffer, index=False)
        col2.download_button("📗 Excel", data=excel_buffer.getvalue(), file_name='relatorio.xlsx', mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

        # Word
        doc = Document()
        doc.add_heading("Relatório de Veículos Elétricos", 0)
        table = doc.add_table(rows=1, cols=len(df.columns))
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(df.columns):
            hdr_cells[i].text = col_name
        for row in df.itertuples(index=False):
            row_cells = table.add_row().cells
            for i, item in enumerate(row):
                row_cells[i].text = str(item)
        word_buffer = BytesIO()
        doc.save(word_buffer)
        col3.download_button("📘 Word", data=word_buffer.getvalue(), file_name="relatorio.docx")

        # PDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=10)
        line_height = pdf.font_size * 2.5
        col_width = pdf.epw / len(df.columns)
        for col_name in df.columns:
            pdf.cell(col_width, line_height, col_name[:20], border=1)
        pdf.ln(line_height)
        for row in df.itertuples(index=False):
            for item in row:
                pdf.cell(col_width, line_height, str(item)[:20], border=1)
            pdf.ln(line_height)
        pdf_buffer = BytesIO()
        pdf.output(pdf_buffer)
        col4.download_button("📕 PDF", data=pdf_buffer.getvalue(), file_name="relatorio.pdf")

    else:
        st.warning("⚠️ Nenhum resultado encontrado para os filtros selecionados.")
else:
    st.info("🧭 Selecione os filtros desejados e clique em **Pesquisar** para exibir os resultados.")
