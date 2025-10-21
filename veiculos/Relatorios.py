import streamlit as st
import pandas as pd
import mysql.connector
from mysql.connector import Error
from io import BytesIO
from docx import Document
from fpdf import FPDF

# ------------------- CONFIGURAÇÃO -------------------
st.set_page_config(
    page_title="Relatórios - Veículos Elétricos",
    layout="wide"
)

# ------------------- ESTILOS -------------------
st.markdown("""
    <style>
        h1, .stTitle {
            font-size: 2.8em !important;
            color: #1a5130;
            text-align: center;
            margin-bottom: 15px;
        }
        h2, .stSubheader {
            font-size: 1.8em !important;
            color: #1a5130;
            margin-bottom: 10px;
        }
        .kpi-card {
            background-color: #EAF6EA;
            padding: 20px;
            border-radius: 10px;
            text-align: center;
            margin-bottom: 20px;
        }
        .kpi-value {
            font-size: 2.4em;
            font-weight: 700;
            color: #1a5130;
        }
        .kpi-label {
            font-size: 1.2em;
            color: #4a7856;
        }
        .card {
            background-color: #F2F8F2;
            padding: 20px;
            border-radius: 10px;
            margin-bottom: 20px;
        }
        .card-title {
            font-size: 1.6em;
            font-weight: 700;
            color: #1a5130;
            margin-bottom: 10px;
        }
        .card-content {
            font-size: 1.1em;
            color: #333333;
        }
        .dataframe tbody tr:nth-child(even) {
            background-color: #f3f9f3;
        }
        .dataframe thead th {
            background-color: #dff0d8;
            color: #1a5130;
        }
        .stButton button {
            border-radius: 8px;
            height: 40px;
            font-weight: 600;
        }
    </style>
""", unsafe_allow_html=True)

# ------------------- TÍTULO -------------------
st.title("📊 Relatórios de Veículos Elétricos")
st.markdown(
    '<p style="text-align:center; font-size:1.4em; color:#4a7856;">Filtragem e exportação de dados da frota de veículos elétricos no Brasil</p>',
    unsafe_allow_html=True
)

# ------------------- CONEXÃO SEGURA -------------------
@st.cache_resource
def conectar():
    """Cria conexão segura com o banco Aiven usando st.secrets e SSL."""
    try:
        conexao = mysql.connector.connect(
            host=st.secrets["MYSQL_HOST"],
            user=st.secrets["MYSQL_USER"],
            password=st.secrets["MYSQL_PASSWORD"],
            database=st.secrets["MYSQL_DATABASE"],
            port=st.secrets["MYSQL_PORT"],
            ssl_ca=st.secrets["SSL_CA_PATH"]
        )
        return conexao
    except Error as e:
        st.error(f"Erro ao conectar ao banco de dados: {e}")
        return None

conn = conectar()

# ------------------- FUNÇÕES AUXILIARES -------------------
def adicionar_opcao_todos(lista):
    return ["Todos"] + lista

def sql_in_list(int_list):
    return "(" + ",".join(str(x) for x in int_list) + ")"

# ------------------- FILTROS -------------------
if conn:
    # Região
    regioes_df = pd.read_sql("SELECT DISTINCT regiao, id_regiao FROM regiao ORDER BY regiao", conn)
    regioes = adicionar_opcao_todos(regioes_df['regiao'].tolist())
    regioes_sel = st.sidebar.multiselect("Regiões", regioes)

    # Estados
    if regioes_sel and "Todos" not in regioes_sel:
        ids_regioes = regioes_df[regioes_df['regiao'].isin(regioes_sel)]['id_regiao'].tolist()
        estados_df = pd.read_sql(f"SELECT uf, id_estado FROM estado WHERE id_regiao IN {sql_in_list(ids_regioes)} ORDER BY uf", conn)
    else:
        estados_df = pd.read_sql("SELECT uf, id_estado FROM estado ORDER BY uf", conn)

    estados = adicionar_opcao_todos(estados_df['uf'].tolist())
    estados_sel = st.sidebar.multiselect("Estados", estados)

    # Cidades
    if estados_sel and "Todos" not in estados_sel:
        ids_estados = estados_df[estados_df['uf'].isin(estados_sel)]['id_estado'].tolist()
        cidades_df = pd.read_sql(f"SELECT cidade, id_cidade FROM cidade WHERE id_estado IN {sql_in_list(ids_estados)} ORDER BY cidade", conn)
    else:
        cidades_df = pd.read_sql("SELECT cidade, id_cidade FROM cidade ORDER BY cidade", conn)

    cidades = adicionar_opcao_todos(cidades_df['cidade'].tolist())
    cidades_sel = st.sidebar.multiselect("Cidades", cidades)

    # Ano
    anos_df = pd.read_sql("SELECT DISTINCT ano_inicial AS ano FROM cidade_tipo_modelo ORDER BY ano_inicial", conn)
    anos = adicionar_opcao_todos(anos_df['ano'].astype(str).tolist())
    anos_sel = st.sidebar.multiselect("Ano", anos)

    # Tecnologia
    tecnologias_df = pd.read_sql("SELECT DISTINCT tecnologia, id_tecnologia FROM tecnologia ORDER BY tecnologia", conn)
    tecnologias = adicionar_opcao_todos(tecnologias_df['tecnologia'].tolist())
    tecnologias_sel = st.sidebar.multiselect("Tecnologia", tecnologias)

    # Classificação
    classificacao_df = pd.read_sql("SELECT DISTINCT classificacao, id_classificacao FROM classificacao_veiculo ORDER BY classificacao", conn)
    classificacoes = adicionar_opcao_todos(classificacao_df['classificacao'].tolist())
    classificacoes_sel = st.sidebar.multiselect("Classificação", classificacoes)

    # ------------------- BOTÕES -------------------
    st.sidebar.markdown("---")
    col_btn1, col_btn2 = st.sidebar.columns(2)
    pesquisar = col_btn1.button("Pesquisar", use_container_width=True)
    limpar = col_btn2.button("Limpar Filtros", use_container_width=True)

    if limpar:
        st.session_state.clear()
        st.experimental_rerun()

    # ------------------- CONSULTA -------------------
    if pesquisar:
        query = """
            SELECT m.modelo, SUM(ctm.quantidade) AS quantidade, t.tecnologia, cv.classificacao
            FROM cidade_tipo_modelo ctm
            JOIN modelo m ON ctm.id_modelo = m.id_modelo
            JOIN tecnologia t ON ctm.id_tecnologia = t.id_tecnologia
            JOIN classificacao_veiculo cv ON ctm.id_classificacao = cv.id_classificacao
            JOIN cidade c ON ctm.id_cidade = c.id_cidade
            JOIN estado e ON c.id_estado = e.id_estado
            JOIN regiao r ON e.id_regiao = r.id_regiao
            WHERE ctm.quantidade > 0
        """
        if regioes_sel and "Todos" not in regioes_sel:
            ids = regioes_df[regioes_df["regiao"].isin(regioes_sel)]["id_regiao"].tolist()
            query += f" AND r.id_regiao IN {sql_in_list(ids)}"
        if estados_sel and "Todos" not in estados_sel:
            ids = estados_df[estados_df["uf"].isin(estados_sel)]["id_estado"].tolist()
            query += f" AND e.id_estado IN {sql_in_list(ids)}"
        if cidades_sel and "Todos" not in cidades_sel:
            ids = cidades_df[cidades_df["cidade"].isin(cidades_sel)]["id_cidade"].tolist()
            query += f" AND c.id_cidade IN {sql_in_list(ids)}"
        if anos_sel and "Todos" not in anos_sel:
            anos_n = [int(a) for a in anos_sel if a != "Todos"]
            query += f" AND (ctm.ano_inicial IN {sql_in_list(anos_n)} OR ctm.ano_final IN {sql_in_list(anos_n)})"
        if tecnologias_sel and "Todos" not in tecnologias_sel:
            ids = tecnologias_df[tecnologias_df["tecnologia"].isin(tecnologias_sel)]["id_tecnologia"].tolist()
            query += f" AND t.id_tecnologia IN {sql_in_list(ids)}"
        if classificacoes_sel and "Todos" not in classificacoes_sel:
            ids = classificacao_df[classificacao_df["classificacao"].isin(classificacoes_sel)]["id_classificacao"].tolist()
            query += f" AND cv.id_classificacao IN {sql_in_list(ids)}"
        query += " GROUP BY m.modelo, t.tecnologia, cv.classificacao ORDER BY m.modelo"

        df = pd.read_sql(query, conn)

        if not df.empty:
            total_veiculos = df['quantidade'].sum()
            filtros_txt = []
            for nome, valor in {
                "Região": regioes_sel,
                "Estado": estados_sel,
                "Cidade": cidades_sel,
                "Ano": anos_sel,
                "Tecnologia": tecnologias_sel,
                "Classificação": classificacoes_sel
            }.items():
                if valor:
                    filtros_txt.append(f"{nome}: {', '.join(map(str, valor))}")

            st.subheader(f"Resultados para a pesquisa: {' | '.join(filtros_txt) if filtros_txt else '(sem filtros)'}")

            # KPIs
            col1, col2 = st.columns(2)
            col1.markdown(
                f'<div class="kpi-card"><div class="kpi-value">{df["modelo"].nunique()}</div><div class="kpi-label">Modelos únicos</div></div>',
                unsafe_allow_html=True
            )
            col2.markdown(
                f'<div class="kpi-card"><div class="kpi-value">{total_veiculos}</div><div class="kpi-label">Total de veículos</div></div>',
                unsafe_allow_html=True
            )

            # Tabela
            st.markdown('<div class="card"><p class="card-title">Tabela de Resultados</p></div>', unsafe_allow_html=True)
            st.dataframe(df, use_container_width=True)

            # Botões de download
            col1, col2, col3, col4 = st.columns(4)
            csv = df.to_csv(index=False).encode('utf-8')
            col1.download_button("📥 CSV", data=csv, file_name='relatorio.csv', mime='text/csv')

            excel_buffer = BytesIO()
            df.to_excel(excel_buffer, index=False)
            col2.download_button("📗 Excel", data=excel_buffer.getvalue(),
                                 file_name='relatorio.xlsx',
                                 mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

            doc = Document()
            doc.add_heading("Relatório de Veículos Elétricos", 0)
            table = doc.add_table(rows=1, cols=len(df.columns))
            hdr_cells = table.rows[0].cells
            for i, col_name in enumerate(df.columns):
                hdr_cells[i].text = col_name
            for row in df.itertuples(index=False):
                row_cells = table.add_row().cells
                for i, item in enumerate(row):
                    row_cells[i].text = str(item)
            word_buffer = BytesIO()
            doc.save(word_buffer)
            col3.download_button("📘 Word", data=word_buffer.getvalue(), file_name="relatorio.docx")

            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=10)
            line_height = pdf.font_size * 2.5
            col_width = pdf.epw / len(df.columns)
            for col_name in df.columns:
                pdf.cell(col_width, line_height, col_name[:20], border=1)
            pdf.ln(line_height)
            for row in df.itertuples(index=False):
                for item in row:
                    pdf.cell(col_width, line_height, str(item)[:20], border=1)
                pdf.ln(line_height)
            pdf_buffer = BytesIO()
            pdf.output(pdf_buffer)
            col4.download_button("📕 PDF", data=pdf_buffer.getvalue(), file_name="relatorio.pdf")

        else:
            st.warning("Nenhum resultado encontrado para os filtros selecionados.")
    else:
        st.info("Selecione os filtros desejados e clique em Pesquisar para exibir os resultados.")
else:
    st.error("❌ Falha na conexão com o banco de dados.")
